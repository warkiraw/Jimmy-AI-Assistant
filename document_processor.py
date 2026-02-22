"""
Модуль обработки и парсинга документов (PDF, Word, Excel).
"""
import os
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pandas as pd
from langchain_text_splitters import RecursiveCharacterTextSplitter


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx"}
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))


def extract_text_from_pdf(file_path: str) -> str:
    """Извлекает текст из PDF-файла с помощью PyMuPDF."""
    text_parts = []
    doc = fitz.open(file_path)
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Извлекает текст из Word-документа (.docx)."""
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            tables_text.append(row_text)
    result = "\n".join(paragraphs)
    if tables_text:
        result += "\n\n" + "\n".join(tables_text)
    return result


def dataframe_to_markdown_table(df: pd.DataFrame) -> str:
    """Конвертирует DataFrame в Markdown-таблицу."""
    if df.empty:
        return ""
    header = "| " + " | ".join(str(col) for col in df.columns) + " |"
    separator = "| " + " | ".join("---" for _ in df.columns) + " |"
    rows = []
    for _, row in df.iterrows():
        rows.append("| " + " | ".join(str(v) for v in row.values) + " |")
    return "\n".join([header, separator] + rows)


def extract_text_from_excel(file_path: str) -> str:
    """Извлекает данные из Excel-файла в формате Markdown-таблиц."""
    excel_file = pd.ExcelFile(file_path)
    parts = []
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
        if not df.empty:
            parts.append(f"## Лист: {sheet_name}\n\n{dataframe_to_markdown_table(df)}")
    return "\n\n".join(parts)


def extract_text(file_path: str) -> str:
    """
    Определяет тип файла по расширению и извлекает текст.
    Поддерживаются: .pdf, .docx, .xlsx
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    if ext == ".xlsx":
        return extract_text_from_excel(file_path)
    raise ValueError(f"Неподдерживаемый формат файла: {ext}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Разбивает текст на чанки с помощью RecursiveCharacterTextSplitter.
    """
    if not text or not text.strip():
        return []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_text(text.strip())


def process_document(file_path: str) -> list[tuple[str, str]]:
    """
    Обрабатывает документ: извлекает текст и разбивает на чанки.
    Возвращает список кортежей (chunk_text, source_filename).
    """
    filename = os.path.basename(file_path)
    text = extract_text(file_path)
    chunks = chunk_text(text)
    return [(chunk, filename) for chunk in chunks]
