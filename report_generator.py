"""
Генерация DOCX-отчётов из текста, сгенерированного LLM.
"""
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(text: str, title: str, output_path: Path) -> None:
    """
    Создаёт DOCX-документ: заголовок, дата, текст (разбитый по абзацам и подзаголовкам).
    """
    doc = Document()
    # Заголовок 1 уровня
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Дата генерации
    doc.add_paragraph(datetime.now().strftime("%d.%m.%Y"), style="Intense Quote")
    doc.add_paragraph()
    # Разбор текста: параграфы по \n\n, заголовки по # ## ###
    blocks = re.split(r"\n\s*\n", text.strip())
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        first = lines[0].strip()
        if first.startswith("### "):
            doc.add_heading(first[4:].strip(), level=2)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
        elif first.startswith("## "):
            doc.add_heading(first[3:].strip(), level=1)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
        elif first.startswith("# "):
            doc.add_heading(first[2:].strip(), level=1)
            rest = "\n".join(lines[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
        else:
            doc.add_paragraph(block.replace("\n", " "))
    doc.save(output_path)


def generate_report_filename() -> str:
    """Возвращает уникальное имя файла отчёта."""
    return f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
